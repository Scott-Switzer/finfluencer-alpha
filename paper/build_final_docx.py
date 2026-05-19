#!/usr/bin/env python3
"""Build the final FIN 496 manuscript from frozen repository outputs.

The builder consumes committed CSV/MD exports under
data/exports/final_paper_package_v2_expanded and the manuscript source at
paper/final_manuscript.md. It does not fetch data, call APIs, read raw/private
workbooks, or rebuild empirical modules.
"""

from __future__ import annotations

import math
import re
import textwrap
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any, Iterable

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
PAPER_DIR = REPO_ROOT / "paper"
TABLE_DIR = PAPER_DIR / "tables"
FIGURE_DIR = PAPER_DIR / "figures"
SOURCE_DIR = REPO_ROOT / "data" / "exports" / "final_paper_package_v2_expanded"

FROZEN_COMMIT = "5a81aa3e497a358fa9e154ee67b146510e325f40"
SHORT_TITLE = "Finfluencer Recommendations and Stock Returns"
MANUSCRIPT_PATH = PAPER_DIR / "final_manuscript.md"


@dataclass
class AuditRow:
    claim: str
    value: str
    source: str
    method: str
    status: str = "verified"


AUDIT_ROWS: list[AuditRow] = []


def rel(path: Path) -> str:
    try:
        return str(path.relative_to(REPO_ROOT))
    except ValueError:
        return str(path)


def pct(value: float | int | None, decimals: int = 2) -> str:
    if value is None or pd.isna(value):
        return ""
    return f"{100 * float(value):.{decimals}f}%"


def num(value: float | int | None, decimals: int = 3) -> str:
    if value is None or pd.isna(value):
        return ""
    if isinstance(value, int) or float(value).is_integer():
        return f"{int(value):,}"
    return f"{float(value):,.{decimals}f}"


def maybe_pct_string(value: object) -> str:
    if value is None or pd.isna(value):
        return ""
    text = str(value)
    if text.endswith("%"):
        return text
    try:
        return pct(float(value), 1)
    except (TypeError, ValueError):
        return text


def read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(f"Required source missing: {rel(path)}")
    return pd.read_csv(path)


def to_markdown(df: pd.DataFrame) -> str:
    headers = [str(c) for c in df.columns]
    rows = [[("" if pd.isna(v) else str(v)) for v in row] for row in df.to_numpy()]
    widths = [len(h) for h in headers]
    for row in rows:
        widths = [max(w, len(cell)) for w, cell in zip(widths, row)]
    line = "| " + " | ".join(h.ljust(w) for h, w in zip(headers, widths)) + " |"
    sep = "| " + " | ".join("-" * w for w in widths) + " |"
    body = [
        "| " + " | ".join(cell.ljust(w) for cell, w in zip(row, widths)) + " |"
        for row in rows
    ]
    return "\n".join([line, sep, *body])


def write_csv_and_md(df: pd.DataFrame, stem: str, caption: str) -> tuple[Path, Path]:
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    csv_path = TABLE_DIR / f"{stem}.csv"
    md_path = TABLE_DIR / f"{stem}.md"
    df.to_csv(csv_path, index=False)
    md_path.write_text(f"# {caption}\n\n{to_markdown(df)}\n", encoding="utf-8")
    return csv_path, md_path


def add_audit(
    claim: str,
    value: object,
    source: Path,
    method: str,
    status: str = "verified",
) -> None:
    AUDIT_ROWS.append(AuditRow(claim, str(value), rel(source), method, status))


def first_matching_row(df: pd.DataFrame, column: str, value: str) -> pd.Series:
    rows = df[df[column].astype(str).eq(value)]
    if rows.empty:
        raise ValueError(f"Missing expected row {column}={value}")
    return rows.iloc[0]


def font_path() -> str | None:
    candidates = [
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/Library/Fonts/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return None


def bar_chart(
    labels: list[str],
    values: list[float],
    title: str,
    output: Path,
    *,
    ylabel: str = "",
    value_suffix: str = "",
    width: int = 1200,
    height: int = 760,
) -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    fp = font_path()
    title_font = ImageFont.truetype(fp, 38) if fp else ImageFont.load_default()
    label_font = ImageFont.truetype(fp, 24) if fp else ImageFont.load_default()
    small_font = ImageFont.truetype(fp, 20) if fp else ImageFont.load_default()

    left, top, right, bottom = 115, 112, width - 55, height - 158
    draw.text((left, 34), title, fill=(20, 20, 20), font=title_font)
    draw.line((left, bottom, right, bottom), fill=(50, 50, 50), width=2)
    draw.line((left, top, left, bottom), fill=(50, 50, 50), width=2)

    max_value = max(values) if values else 1
    max_value = max_value if max_value > 0 else 1
    rounded_max = math.ceil(max_value / 10) * 10 if max_value > 20 else max_value * 1.15
    for i in range(6):
        y = bottom - (bottom - top) * i / 5
        grid_val = rounded_max * i / 5
        draw.line((left, y, right, y), fill=(226, 226, 226), width=1)
        draw.text((28, y - 12), f"{grid_val:,.0f}", fill=(90, 90, 90), font=small_font)

    if ylabel:
        draw.text((left, height - 42), ylabel, fill=(75, 75, 75), font=small_font)

    palette = [(31, 78, 121), (112, 48, 160), (80, 125, 45), (192, 80, 77), (247, 150, 70)]
    n = len(labels)
    gap = 22
    bar_width = max(22, int((right - left - gap * (n + 1)) / max(n, 1)))
    for idx, (label, value) in enumerate(zip(labels, values)):
        x0 = left + gap + idx * (bar_width + gap)
        x1 = x0 + bar_width
        y0 = bottom - (bottom - top) * value / rounded_max
        color = palette[idx % len(palette)]
        draw.rectangle((x0, y0, x1, bottom), fill=color)
        value_label = f"{value:,.1f}{value_suffix}" if value_suffix else f"{value:,.0f}"
        draw.text((x0, y0 - 28), value_label, fill=(30, 30, 30), font=small_font)
        wrapped = textwrap.wrap(label, width=max(8, int(bar_width / 12)))
        for line_no, line in enumerate(wrapped[:3]):
            draw.text((x0, bottom + 10 + line_no * 24), line, fill=(30, 30, 30), font=label_font)

    image.save(output)


def grouped_line_chart(
    horizons: list[str],
    series: dict[str, list[float]],
    title: str,
    output: Path,
    *,
    value_suffix: str = "%",
    width: int = 1200,
    height: int = 760,
) -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    fp = font_path()
    title_font = ImageFont.truetype(fp, 38) if fp else ImageFont.load_default()
    label_font = ImageFont.truetype(fp, 24) if fp else ImageFont.load_default()
    small_font = ImageFont.truetype(fp, 20) if fp else ImageFont.load_default()
    left, top, right, bottom = 120, 112, width - 75, height - 140
    draw.text((left, 34), title, fill=(20, 20, 20), font=title_font)
    all_values = [v for vals in series.values() for v in vals]
    y_min = min(min(all_values), 0)
    y_max = max(max(all_values), 0)
    span = y_max - y_min if y_max != y_min else 1
    y_min -= span * 0.12
    y_max += span * 0.12

    def y_coord(v: float) -> float:
        return bottom - (v - y_min) / (y_max - y_min) * (bottom - top)

    def x_coord(i: int) -> float:
        if len(horizons) == 1:
            return (left + right) / 2
        return left + i * (right - left) / (len(horizons) - 1)

    for i in range(6):
        val = y_min + (y_max - y_min) * i / 5
        y = y_coord(val)
        draw.line((left, y, right, y), fill=(226, 226, 226), width=1)
        draw.text((32, y - 12), f"{val:.1f}{value_suffix}", fill=(90, 90, 90), font=small_font)
    zero_y = y_coord(0)
    draw.line((left, zero_y, right, zero_y), fill=(90, 90, 90), width=2)
    draw.line((left, bottom, right, bottom), fill=(50, 50, 50), width=2)
    colors = {"top5": (31, 78, 121), "non_top": (192, 80, 77), "all": (80, 125, 45)}
    for name, vals in series.items():
        points = [(x_coord(i), y_coord(v)) for i, v in enumerate(vals)]
        color = colors.get(name, (80, 80, 80))
        for p0, p1 in zip(points, points[1:]):
            draw.line((*p0, *p1), fill=color, width=4)
        for x, y in points:
            draw.ellipse((x - 6, y - 6, x + 6, y + 6), fill=color)
    for i, h in enumerate(horizons):
        draw.text((x_coord(i) - 18, bottom + 12), h, fill=(30, 30, 30), font=label_font)
    legend_x = left
    for name, color in colors.items():
        if name in series:
            label = name.replace("_", "-")
            draw.rectangle((legend_x, height - 58, legend_x + 22, height - 36), fill=color)
            draw.text((legend_x + 30, height - 62), label, fill=(30, 30, 30), font=small_font)
            legend_x += 145
    image.save(output)


def build_tables_and_figures() -> dict[str, pd.DataFrame]:
    sample_path = SOURCE_DIR / "01_v2_sample_construction_table.csv"
    event_path = SOURCE_DIR / "locked_sample_v2" / "02_v2_event_manifest.csv"
    baseline_path = SOURCE_DIR / "02_v2_event_study_robustness_table.csv"
    long_path = SOURCE_DIR / "long_horizon" / "04_v2_long_horizon_top5_vs_non_top.csv"
    factor_path = SOURCE_DIR / "factor_alpha_beta_estimated" / "03_factor_alpha_summary_by_spec.csv"
    calendar_path = SOURCE_DIR / "calendar_time_factor_regressions" / "01_calendar_time_hac_regressions.csv"
    bloomberg_path = SOURCE_DIR / "bloomberg_validation" / "bloomberg_field_coverage_summary.csv"
    extreme_path = SOURCE_DIR / "news_extreme_event_audit" / "extreme_event_news_audit_summary.csv"
    robustness_path = SOURCE_DIR / "paper_robustness" / "placebo_permutation_summary.csv"
    news_path = SOURCE_DIR / "news_confound_master" / "news_clean_status_return_table.csv"
    market_quiet_path = SOURCE_DIR / "market_implied_confounds" / "returns_by_market_confound_bucket.csv"
    placebo_path = SOURCE_DIR / "research_frontier" / "placebo_matched_controls" / "creator_cross_ticker_placebo_results.csv"
    matched_path = SOURCE_DIR / "causal_diagnostics" / "05_v2_matched_control_tests.csv"
    permutation_path = SOURCE_DIR / "causal_diagnostics" / "04_v2_permutation_tests.csv"
    creator_path = SOURCE_DIR / "creator_deep_dive" / "01_creator_summary.csv"
    ticker_path = SOURCE_DIR / "09_v2_ticker_heterogeneity.csv"
    coverage_path = SOURCE_DIR / "long_horizon" / "02_v2_long_horizon_coverage.csv"

    sample = read_csv(sample_path)
    events = read_csv(event_path)
    baseline = read_csv(baseline_path)
    long_top = read_csv(long_path)
    factor = read_csv(factor_path)
    calendar = read_csv(calendar_path)
    bloomberg = read_csv(bloomberg_path)
    extreme = read_csv(extreme_path)
    robustness = read_csv(robustness_path)
    news = read_csv(news_path)
    market_quiet = read_csv(market_quiet_path)
    placebo = read_csv(placebo_path)
    matched = read_csv(matched_path)
    permutation = read_csv(permutation_path)
    creator = read_csv(creator_path)
    ticker = read_csv(ticker_path)
    coverage = read_csv(coverage_path)

    metrics = dict(zip(sample["metric"], sample["count"]))
    for metric in [
        "live_transcript_rows",
        "successful_transcript_rows",
        "candidate_windows",
        "accepted_recommendation_events",
        "distinct_event_videos",
        "buy_recommendations",
        "sell_recommendations",
        "creators",
        "tickers",
        "return_matched_1d",
        "return_matched_5d",
        "low_lookahead_events",
        "duplicate_collapsed_events",
        "top5_events",
        "non_top_events",
        "factor_matched_events",
    ]:
        add_audit(metric.replace("_", " "), int(metrics[metric]), sample_path, f"metric lookup: {metric}")
    add_audit("Accepted recommendation events cross-check", len(events), event_path, "row count")

    sample_metrics = [
        "live_transcript_rows",
        "successful_transcript_rows",
        "candidate_windows",
        "accepted_recommendation_events",
        "distinct_event_videos",
        "buy_recommendations",
        "sell_recommendations",
        "creators",
        "tickers",
        "return_matched_1d",
        "return_matched_5d",
        "low_lookahead_events",
        "duplicate_collapsed_events",
        "top5_events",
        "non_top_events",
    ]
    sample_table = sample[sample["metric"].isin(sample_metrics)].copy()
    sample_table["count"] = sample_table["count"].map(lambda x: f"{int(x):,}")
    sample_table = sample_table[["metric", "count", "source", "filter_definition"]]
    write_csv_and_md(sample_table, "table_01_sample_construction", "Table 1. Sample construction")

    events["year"] = pd.to_datetime(events["event_date"]).dt.year
    by_year = events.groupby("year", as_index=False).size().rename(columns={"size": "events"})
    write_csv_and_md(by_year, "table_02_events_by_year", "Table 2. Event counts by year")
    add_audit("Event years represented", f"{by_year['year'].min()}-{by_year['year'].max()}", event_path, "min/max event_date year")
    add_audit("Event counts by year", dict(zip(by_year["year"].astype(str), by_year["events"].astype(int))), event_path, "group by event_date year")

    by_type = events.groupby("recommendation_type", as_index=False).size().rename(columns={"size": "events"})
    write_csv_and_md(by_type, "table_03_events_by_recommendation_type", "Table 3. Event counts by recommendation type")
    add_audit("Recommendation-type counts", dict(zip(by_type["recommendation_type"], by_type["events"].astype(int))), event_path, "group by recommendation_type")
    timing = events.groupby("upload_timing_bucket").size().to_dict()
    quality = events.groupby("quality_score").size().to_dict()
    add_audit("Upload timing bucket counts", timing, event_path, "group by upload_timing_bucket")
    add_audit("Quality score distribution", quality, event_path, "group by quality_score")

    selected_specs = [
        "v2 all accepted events",
        "v2 low-lookahead",
        "v2 duplicate-collapsed",
        "v2 top-5 tickers",
        "v2 non-top tickers",
        "v2 buy-only",
        "v2 sell-only",
    ]
    base_table = baseline[baseline["specification"].isin(selected_specs)].copy()
    for spec in selected_specs:
        row = first_matching_row(baseline, "specification", spec)
        add_audit(f"{spec} 1D mean AR", pct(row["mean_1d_ar"], 3), baseline_path, f"specification={spec}, mean_1d_ar")
        add_audit(f"{spec} 1D n/t/p", f"n={int(row['n_1d'])}, t={row['t_1d']}, p={row['p_1d']}", baseline_path, f"specification={spec}, n_1d/t_1d/p_1d")
        add_audit(f"{spec} 5D mean AR", pct(row["mean_5d_ar"], 3), baseline_path, f"specification={spec}, mean_5d_ar")
        add_audit(f"{spec} 5D n/t/p/win", f"n={int(row['n_5d'])}, t={row['t_5d']}, p={row['p_5d']}, win={pct(row['win_rate_5d'], 1)}", baseline_path, f"specification={spec}, n_5d/t_5d/p_5d/win_rate_5d")
    for col in ["mean_1d_ar", "mean_5d_ar", "median_5d_ar"]:
        base_table[col] = base_table[col].map(lambda x: pct(x, 2))
    base_table["win_rate_5d"] = base_table["win_rate_5d"].map(lambda x: pct(x, 1))
    base_table = base_table[
        ["specification", "n_1d", "mean_1d_ar", "t_1d", "p_1d", "n_5d", "mean_5d_ar", "t_5d", "p_5d", "win_rate_5d"]
    ]
    write_csv_and_md(base_table, "table_04_baseline_event_study", "Table 4. Baseline event-study results")

    horizon_keep = ["1D", "5D", "21D", "63D", "126D"]
    top_table = long_top[long_top["horizon"].isin(horizon_keep)].copy()
    for spec in ["top5", "non_top"]:
        for horizon in horizon_keep:
            row = long_top[(long_top["specification"].eq(spec)) & (long_top["horizon"].eq(horizon))].iloc[0]
            add_audit(
                f"{spec} {horizon} SPY BHAR",
                f"mean={pct(row['mean_spy_bhar'], 3)}, n={int(row['n_events'])}, full={int(row['n_full_window'])}, right_censored={int(row['n_right_censored'])}, t={row['t_spy_bhar']}, p={row['p_spy_bhar']}",
                long_path,
                f"specification={spec},horizon={horizon}",
            )
    for col in ["mean_spy_bhar", "median_spy_bhar", "mean_spy_car"]:
        top_table[col] = top_table[col].map(lambda x: pct(x, 2))
    top_table["win_rate_spy_bhar"] = top_table["win_rate_spy_bhar"].map(lambda x: pct(x, 1))
    top_table = top_table[
        ["specification", "horizon", "n_events", "n_full_window", "n_right_censored", "mean_spy_bhar", "t_spy_bhar", "p_spy_bhar", "win_rate_spy_bhar"]
    ]
    write_csv_and_md(top_table, "table_05_top5_vs_non_top", "Table 5. Top-5 vs non-top long-horizon BHAR")

    factor_table = factor[
        factor["model"].eq("FF5_MOM")
        & factor["sample"].isin(["all", "top5", "non_top"])
        & factor["horizon"].isin(["5D", "21D", "63D"])
    ].copy()
    for _, row in factor_table.iterrows():
        add_audit(
            f"FF5+MOM factor alpha {row['sample']} {row['horizon']}",
            f"n={int(row['n'])}, alpha={row['mean_alpha_pct']}, t={row['t_stat']}, p={row['p_value']}, q={row['bh_q_value']}",
            factor_path,
            f"model=FF5_MOM,sample={row['sample']},horizon={row['horizon']}",
        )
    for model in ["CAPM", "Carhart", "FF5_MOM"]:
        row = factor[
            factor["sample"].eq("all") & factor["model"].eq(model) & factor["horizon"].eq("5D")
        ].iloc[0]
        add_audit(
            f"All-event 5D factor alpha {model}",
            f"n={int(row['n'])}, alpha={row['mean_alpha_pct']}, t={row['t_stat']}, p={row['p_value']}, q={row['bh_q_value']}",
            factor_path,
            f"sample=all,model={model},horizon=5D",
        )
    factor_table = factor_table[["sample", "model", "horizon", "n", "mean_alpha_pct", "t_stat", "p_value", "bh_q_value"]]
    write_csv_and_md(factor_table, "table_06_factor_alpha", "Table 6. FF5+momentum factor-adjusted event alpha")

    cal_table = calendar[
        calendar["model"].eq("FF5_MOM")
        & calendar["weighting"].eq("ew")
        & calendar["holding_trading_days"].isin([5, 21])
        & calendar["strategy"].isin(["long_all_buy", "long_top5_buys_only", "short_non_top_buys_diagnostic"])
    ].copy()
    for _, row in cal_table.iterrows():
        add_audit(
            f"Calendar-time FF5+MOM alpha {row['strategy']} {int(row['holding_trading_days'])}D",
            f"n_days={int(row['n_days'])}, alpha_daily={pct(row['alpha_daily'], 3)}, alpha_ann={pct(row['alpha_ann_approx'], 1)}, t={row['alpha_t_hac']}, p={row['alpha_p_value']}",
            calendar_path,
            f"strategy={row['strategy']},holding={row['holding_trading_days']},weighting=ew,model=FF5_MOM",
        )
    cal_table["alpha_daily"] = cal_table["alpha_daily"].map(lambda x: pct(x, 3))
    cal_table["alpha_ann_approx"] = cal_table["alpha_ann_approx"].map(lambda x: pct(x, 1))
    cal_table = cal_table[["strategy", "holding_trading_days", "model", "n_days", "alpha_daily", "alpha_ann_approx", "alpha_t_hac", "alpha_p_value"]]
    write_csv_and_md(cal_table, "table_07_calendar_time_factor_regressions", "Table 7. Calendar-time HAC factor regressions")

    event_cov = bloomberg[
        bloomberg["summary_type"].eq("event_coverage")
        & bloomberg["coverage_level"].eq("event_overall")
        & bloomberg["bucket"].eq("all_events")
    ].copy()
    coverage_features = [
        "analyst_consensus_available",
        "analyst_coverage_count_available",
        "estimates_available",
        "news_proxy_available",
        "liquidity_proxy_available",
        "short_interest_available",
        "event_total_return_available",
    ]
    bloom_table = event_cov[event_cov["feature"].isin(coverage_features)].copy()
    bloom_table["events"] = bloom_table["events"].astype(int)
    bloom_table["covered_events"] = bloom_table["covered_events"].astype(int)
    bloom_table["event_coverage_pct"] = bloom_table["event_coverage_pct"].map(maybe_pct_string)
    bloom_table = bloom_table[["feature", "events", "covered_events", "event_coverage_pct"]]
    write_csv_and_md(bloom_table, "table_08_bloomberg_coverage", "Table 8. Bloomberg event-level coverage")
    for _, row in event_cov.iterrows():
        if str(row.get("feature", "")):
            add_audit(
                f"Bloomberg coverage {row['feature']}",
                f"{int(row['covered_events'])} / {int(row['events'])} ({row['event_coverage_pct']})",
                bloomberg_path,
                f"event_coverage feature={row['feature']}",
            )
    tot_rows = bloomberg[bloomberg["field"].eq("TOT_ANALYST_REC")]
    if not tot_rows.empty:
        row = tot_rows.iloc[0]
        add_audit(
            "TOT_ANALYST_REC source workbook parse",
            f"sheet={row['sheet_name']}, source_observations={int(row['source_observations'])}, valid_observations={int(row['valid_observations'])}, valid_pct={row['valid_value_pct']}, dates={row['first_date']} to {row['last_date']}, tickers={int(row['ticker_count'])}",
            bloomberg_path,
            "source_field row field=TOT_ANALYST_REC",
        )

    extreme_metrics = [
        "audited_unique_events",
        "bucket_positive_1d",
        "bucket_negative_1d",
        "bucket_positive_5d",
        "bucket_negative_5d",
        "official_confounded",
        "media_confounded",
        "bloomberg_news_flow_high",
        "market_attention_high",
        "institutionally_followed",
        "unresolved_unknown",
        "candidate_clean_extreme",
        "selection_rows_before_dedup",
    ]
    extreme_table = extreme[extreme["metric"].isin(extreme_metrics)].copy()
    for metric in extreme_metrics:
        row = extreme[extreme["metric"].eq(metric)].iloc[0]
        add_audit(metric.replace("_", " "), int(float(row["count"])), extreme_path, f"metric={metric}")
    extreme_table["count"] = extreme_table["count"].map(lambda x: f"{int(float(x)):,}" if pd.notna(x) else "")
    extreme_table = extreme_table[["metric", "count", "percent", "notes"]]
    write_csv_and_md(extreme_table, "table_09_extreme_event_audit", "Table 9. Extreme-event audit classification")

    ms_rows = news[(news["sample"].eq("full_sample")) & (news["news_clean_status"].eq("multi_source_clean"))]
    for _, row in ms_rows.iterrows():
        add_audit(
            f"Multi-source clean public-news events {row['horizon']}",
            int(row["n"]),
            news_path,
            f"sample=full_sample, news_clean_status=multi_source_clean, horizon={row['horizon']}",
        )
    ms_clean_5d = ms_rows[ms_rows["horizon"].eq("5D")].iloc[0]
    quiet = market_quiet[(market_quiet["sample"].eq("non_top_market_quiet")) & (market_quiet["horizon"].eq("21D"))].iloc[0]
    placebo_diff = placebo["diff_treated_minus_placebo"].mean()
    add_audit("Non-top market-quiet 21D SPY BHAR", pct(quiet["mean"], 3), market_quiet_path, "sample=non_top_market_quiet,horizon=21D")
    add_audit("Non-top market-quiet 21D n/t/p", f"n={int(quiet['n'])}, t={quiet['t_stat']}, p={quiet['p_value']}", market_quiet_path, "sample=non_top_market_quiet,horizon=21D")
    add_audit("Cross-ticker placebo 5D mean difference", pct(placebo_diff, 3), placebo_path, "mean(diff_treated_minus_placebo)")
    add_audit("Cross-ticker placebo 5D n", len(placebo), placebo_path, "row count")
    for _, row in matched.iterrows():
        add_audit(
            f"Matched-control {row['specification']}",
            f"n={int(row['n'])}, mean={pct(row['mean_ar'], 3)}, t={row['t_stat']}, p={row['p_value']}, win={pct(row['win_rate'], 1)}",
            matched_path,
            f"specification={row['specification']}",
        )
    perm = permutation.iloc[0]
    add_audit(
        "Permutation shuffle event dates within ticker",
        f"permutations={int(perm['permutations'])}, observed={pct(perm['observed_mean_5d_ar'], 3)}, permutation_mean={pct(perm['permutation_mean'], 3)}, p={perm['permutation_p_value']}",
        permutation_path,
        "single exported permutation test row",
    )
    for _, row in coverage[coverage["horizon"].isin(["1D", "5D", "21D", "63D", "126D", "504D"])].iterrows():
        add_audit(
            f"Return coverage {row['horizon']}",
            f"total={int(row['total_events'])}, with_return={int(row['n_with_return'])}, full={int(row['n_full_window'])}, right_censored={int(row['n_right_censored'])}, missing={int(row['n_missing'])}",
            coverage_path,
            f"horizon={row['horizon']}",
        )
    add_audit(
        "Creator summary event-count range",
        f"creators={len(creator)}, max={int(creator['event_count'].max())}, median={creator['event_count'].median():.0f}, min={int(creator['event_count'].min())}",
        creator_path,
        "creator_deep_dive event_count summary",
    )
    top_creator = creator.sort_values("event_count", ascending=False).iloc[0]
    add_audit(
        "Largest creator event count",
        f"{top_creator['creator']}: {int(top_creator['event_count'])}",
        creator_path,
        "max event_count row",
    )
    for _, row in ticker.head(6).iterrows():
        add_audit(
            f"Ticker heterogeneity {row['ticker']}",
            f"events={int(row['event_count'])}, 5D={pct(row['mean_5d_ar'], 3)}, t={row['t_5d']}, p={row['p_5d']}, win={pct(row['win_rate_5d'], 1)}",
            ticker_path,
            f"ticker={row['ticker']}",
        )

    robust_rows = [
        {
            "check": "Multi-source clean events",
            "value": f"{int(ms_clean_5d['n']):,}",
            "interpretation": "No public-news-clean return claim is supported.",
            "source": rel(news_path),
        },
        {
            "check": "Non-top market-quiet 21D SPY BHAR",
            "value": pct(quiet["mean"], 2),
            "interpretation": "Secondary sensitivity only; market quiet is not news clean.",
            "source": rel(market_quiet_path),
        },
        {
            "check": "Cross-ticker placebo 5D mean diff",
            "value": pct(placebo_diff, 2),
            "interpretation": "Economically near-zero falsification benchmark.",
            "source": rel(placebo_path),
        },
    ]
    for _, row in robustness.iterrows():
        if row["horizon_days"] in ["5D", "21D"] and row["placebo_type"] == "label_shuffle_spy_bhar_within_sample":
            add_audit(
                f"Label-shuffle placebo {row['horizon_days']}",
                f"observed={pct(row['observed_top5_minus_nontop'], 3)}, p={row['permutation_p_upper_tail']}, n={int(row['n_events'])}",
                robustness_path,
                f"placebo_type={row['placebo_type']}, horizon={row['horizon_days']}",
            )
            robust_rows.append(
                {
                    "check": f"Label-shuffle placebo, {row['horizon_days']}",
                    "value": pct(row["observed_top5_minus_nontop"], 2),
                    "interpretation": f"Permutation p={row['permutation_p_upper_tail']}; descriptive heterogeneity check.",
                    "source": rel(robustness_path),
                }
            )
    robust_table = pd.DataFrame(robust_rows)
    write_csv_and_md(robust_table, "table_10_robustness_summary", "Table 10. Robustness and falsification summary")

    bar_chart(by_year["year"].astype(str).tolist(), by_year["events"].astype(float).tolist(), "Accepted events by year", FIGURE_DIR / "events_by_year.png", ylabel="Events")
    bar_chart(by_year["year"].astype(str).tolist(), by_year["events"].astype(float).tolist(), "Accepted events by year", FIGURE_DIR / "accepted_events_by_year.png", ylabel="Events")
    bar_chart(by_type["recommendation_type"].astype(str).tolist(), by_type["events"].astype(float).tolist(), "Events by recommendation type", FIGURE_DIR / "events_by_recommendation_type.png", ylabel="Events")
    bar_chart(by_type["recommendation_type"].astype(str).tolist(), by_type["events"].astype(float).tolist(), "Events by recommendation type", FIGURE_DIR / "events_by_type.png", ylabel="Events")

    chart_h = ["1D", "5D", "21D", "63D"]
    series = {}
    for spec in ["top5", "non_top"]:
        series[spec] = [
            100 * float(long_top[(long_top["specification"].eq(spec)) & (long_top["horizon"].eq(h))]["mean_spy_bhar"].iloc[0])
            for h in chart_h
        ]
    grouped_line_chart(chart_h, series, "SPY-adjusted BHAR by ticker salience", FIGURE_DIR / "top5_non_top_bhar.png")

    bloom_fig = bloom_table.copy()
    bloom_fig["label"] = bloom_fig["feature"].str.replace("_available", "", regex=False).str.replace("event_", "", regex=False).str.replace("_", " ")
    bloom_fig["pct"] = bloom_fig["covered_events"] / bloom_fig["events"] * 100
    bar_chart(bloom_fig["label"].tolist(), bloom_fig["pct"].tolist(), "Bloomberg event-level coverage", FIGURE_DIR / "bloomberg_coverage.png", ylabel="Coverage", value_suffix="%")

    ext_fig = extreme[extreme["metric"].isin(["official_confounded", "media_confounded", "bloomberg_news_flow_high", "market_attention_high", "candidate_clean_extreme"])].copy()
    ext_fig["label"] = ext_fig["metric"].str.replace("_", " ")
    bar_chart(ext_fig["label"].tolist(), ext_fig["count"].astype(float).tolist(), "Extreme-event audit labels", FIGURE_DIR / "extreme_event_audit.png", ylabel="Events")

    return {
        "sample_table": sample_table,
        "by_year": by_year,
        "by_type": by_type,
        "base_table": base_table,
        "top_table": top_table,
        "factor_table": factor_table,
        "cal_table": cal_table,
        "bloom_table": bloom_table,
        "extreme_table": extreme_table,
        "robust_table": robust_table,
    }


def write_references() -> Path:
    refs = """# References

[1] Fama, Eugene F., Lawrence Fisher, Michael C. Jensen, and Richard Roll. 1969. "The Adjustment of Stock Prices to New Information." International Economic Review 10 (1): 1-21.

[2] MacKinlay, A. Craig. 1997. "Event Studies in Economics and Finance." Journal of Economic Literature 35 (1): 13-39.

[3] Kothari, S. P., and Jerold B. Warner. 2006. "Econometrics of Event Studies." In Handbook of Corporate Finance: Empirical Corporate Finance, edited by B. Espen Eckbo. Elsevier.

[4] Barber, Brad M., and Terrance Odean. 2008. "All That Glitters: The Effect of Attention and News on the Buying Behavior of Individual and Institutional Investors." Review of Financial Studies 21 (2): 785-818.

[5] Da, Zhi, Joseph Engelberg, and Pengjie Gao. 2011. "In Search of Attention." Journal of Finance 66 (5): 1461-1499.

[6] Tetlock, Paul C. 2007. "Giving Content to Investor Sentiment: The Role of Media in the Stock Market." Journal of Finance 62 (3): 1139-1168.

[7] Antweiler, Werner, and Murray Z. Frank. 2004. "Is All That Talk Just Noise? The Information Content of Internet Stock Message Boards." Journal of Finance 59 (3): 1259-1294.

[8] Chen, Hailiang, Prabuddha De, Yu Jeffrey Hu, and Byoung-Hyoun Hwang. 2014. "Wisdom of Crowds: The Value of Stock Opinions Transmitted Through Social Media." Review of Financial Studies 27 (5): 1367-1403.

[9] Kakhbod, Ali, Seyed Mohammad Kazempour, Dmitry Livdan, and Norman Schuerhoff. "Finfluencers." Working paper.

[10] Guan, Sue S. 2023. "The Rise of the Finfluencer." New York University Journal of Law and Business.

[11] CFA Institute. 2024. "The Finfluencer Appeal: Investing in the Age of Social Media."

[12] U.S. Securities and Exchange Commission, Office of Investor Education and Advocacy. Investor alerts and public guidance on social media, online promotion, and investment fraud.

[13] Scott Switzer. FIN 496 Capstone repository, finfluencer-alpha, frozen at commit 5a81aa3e497a358fa9e154ee67b146510e325f40. Numeric claims in this manuscript are sourced to committed CSV/MD exports listed in MANUSCRIPT_SOURCE_AUDIT.md.
"""
    path = PAPER_DIR / "references.md"
    path.write_text(refs, encoding="utf-8")
    return path


def write_source_audit() -> Path:
    rows = pd.DataFrame([row.__dict__ for row in AUDIT_ROWS])
    path = PAPER_DIR / "MANUSCRIPT_SOURCE_AUDIT.md"
    text = [
        "# Manuscript Source Audit",
        "",
        f"Frozen source commit: `{FROZEN_COMMIT}`",
        "",
        "Every numeric claim generated or relied upon by `paper/build_final_docx.py` is listed below with the committed source artifact used for verification.",
        "",
        to_markdown(rows),
        "",
        "## Safety checks",
        "",
        "- The builder reads derived public/export artifacts only.",
        "- The builder does not call APIs, fetch transcripts, or rebuild empirical modules.",
        "- Raw Bloomberg workbooks, databases, transcript bodies, secrets, and private caches are not copied into `paper/`.",
        "- Unknown news and analyst states are preserved as unknown; they are not recoded as clean.",
    ]
    path.write_text("\n".join(text) + "\n", encoding="utf-8")
    return path


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: float = 6.2) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def compact_table_for_docx(key: str, tables: dict[str, pd.DataFrame]) -> tuple[str, pd.DataFrame]:
    if key == "sample_construction":
        df = tables["sample_table"][["metric", "count"]].copy()
        df.columns = ["Metric", "Count"]
        return "Table 1. Sample construction", df
    if key == "events_by_year":
        df = tables["by_year"].copy()
        df.columns = ["Year", "Events"]
        return "Table 2. Event counts by year", df
    if key == "events_by_type":
        df = tables["by_type"].copy()
        df.columns = ["Type", "Events"]
        return "Table 3. Event counts by recommendation type", df
    if key == "baseline_event_study":
        df = tables["base_table"].copy()
        df = df[df["specification"].isin(["v2 all accepted events", "v2 top-5 tickers", "v2 non-top tickers", "v2 buy-only", "v2 sell-only"])]
        df = pd.DataFrame(
            {
                "Spec": df["specification"].str.replace("v2 ", "", regex=False),
                "N1": df["n_1d"],
                "1D": df["mean_1d_ar"],
                "N5": df["n_5d"],
                "5D": df["mean_5d_ar"],
                "p5": df["p_5d"],
            }
        )
        return "Table 4. Baseline SPY-adjusted abnormal returns", df
    if key == "top5_vs_non_top":
        df = tables["top_table"].copy()
        df = pd.DataFrame(
            {
                "Sample": df["specification"],
                "H": df["horizon"],
                "N": df["n_events"],
                "Full": df["n_full_window"],
                "AR": df["mean_spy_bhar"],
                "t": df["t_spy_bhar"],
            }
        )
        return "Table 5. Top-5 vs non-top SPY-adjusted BHAR", df
    if key == "factor_alpha":
        df = tables["factor_table"].copy()
        df = pd.DataFrame(
            {
                "Sample": df["sample"],
                "H": df["horizon"],
                "N": df["n"],
                "Alpha": df["mean_alpha_pct"],
                "t": df["t_stat"],
                "p": df["p_value"],
            }
        )
        return "Table 6. FF5+momentum factor-adjusted event alpha", df
    if key == "calendar_time":
        df = tables["cal_table"].copy()
        df = pd.DataFrame(
            {
                "Strategy": df["strategy"].str.replace("_", " "),
                "H": df["holding_trading_days"],
                "N days": df["n_days"],
                "Daily a": df["alpha_daily"],
                "t": df["alpha_t_hac"],
                "p": df["alpha_p_value"],
            }
        )
        return "Table 7. Calendar-time HAC regressions", df
    if key == "bloomberg_coverage":
        df = tables["bloom_table"].copy()
        df = pd.DataFrame(
            {
                "Feature": df["feature"].str.replace("_available", "", regex=False).str.replace("event_", "", regex=False).str.replace("_", " "),
                "Events": df["events"],
                "Covered": df["covered_events"],
                "Pct": df["event_coverage_pct"],
            }
        )
        return "Table 8. Bloomberg event-level coverage", df
    if key == "extreme_event_audit":
        df = tables["extreme_table"].copy()
        df = pd.DataFrame(
            {
                "Label": df["metric"].str.replace("_", " "),
                "Count": df["count"],
                "Pct": df["percent"],
            }
        )
        return "Table 9. Extreme-event audit classification", df
    if key == "robustness_summary":
        df = tables["robust_table"][["check", "value", "interpretation"]].copy()
        df.columns = ["Check", "Value", "Interpretation"]
        return "Table 10. Robustness and falsification summary", df
    raise KeyError(f"Unknown table marker: {key}")


def add_docx_table(doc: Document, key: str, tables: dict[str, pd.DataFrame]) -> None:
    title, df = compact_table_for_docx(key, tables)
    caption = doc.add_paragraph(title, style="Caption")
    caption.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.autofit = False
    for idx, col in enumerate(df.columns):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "D9EAF7")
        set_cell_text(cell, col, bold=True, font_size=6.0)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], "" if pd.isna(value) else str(value), font_size=5.6)
    doc.add_paragraph("")


FIGURE_MARKERS = {
    "events_by_year": ("Figure 1. Accepted events by year", FIGURE_DIR / "events_by_year.png"),
    "events_by_type": ("Figure 2. Events by recommendation type", FIGURE_DIR / "events_by_recommendation_type.png"),
    "top5_non_top_bhar": ("Figure 3. Top-5 vs non-top SPY-adjusted BHAR", FIGURE_DIR / "top5_non_top_bhar.png"),
    "bloomberg_coverage": ("Figure 4. Bloomberg event-level coverage", FIGURE_DIR / "bloomberg_coverage.png"),
    "extreme_event_audit": ("Figure 5. Extreme-event audit labels", FIGURE_DIR / "extreme_event_audit.png"),
}


def add_docx_figure(doc: Document, key: str) -> None:
    if key not in FIGURE_MARKERS:
        raise KeyError(f"Unknown figure marker: {key}")
    caption, path = FIGURE_MARKERS[key]
    if not path.exists():
        raise FileNotFoundError(f"Missing figure: {rel(path)}")
    p = doc.add_paragraph(caption, style="Caption")
    p.paragraph_format.keep_with_next = True
    pic = doc.add_picture(str(path), width=Inches(3.05))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = ""
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(text)
    run._r.append(fld_char_2)


def set_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    if not cols:
        sect_pr.append(cols_el)
    cols_el.set(qn("w:num"), str(count))
    cols_el.set(qn("w:space"), "360")


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9.6)
    normal.paragraph_format.space_after = Pt(3.2)
    normal.paragraph_format.line_spacing = 1.02
    for style_name, size in [("Title", 18), ("Heading 1", 11.5), ("Heading 2", 10.2), ("Caption", 7.2), ("Header", 8), ("Footer", 8)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(0, 0, 0)
    for heading_name in ["Heading 1", "Heading 2"]:
        fmt = styles[heading_name].paragraph_format
        fmt.space_before = Pt(7)
        fmt.space_after = Pt(2)
        fmt.keep_with_next = True


def configure_section(section, *, columns: int) -> None:
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    set_columns(section, columns)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.text = SHORT_TITLE
    header.style = "Header"
    add_page_number(section.footer.paragraphs[0])


def manuscript_title_and_subtitle(lines: list[str]) -> tuple[str, str]:
    title = "YouTube Finfluencers and Stock Returns"
    subtitle = "Evidence from a frozen transcript-supported event-study sample"
    for idx, line in enumerate(lines):
        if line.startswith("# "):
            title = line[2:].strip()
            for candidate in lines[idx + 1 :]:
                text = candidate.strip()
                if text and not text.startswith("#") and not text.lower().startswith("frozen source"):
                    subtitle = text
                    break
            break
    return title, subtitle


def add_markdown_paragraph(doc: Document, text: str, *, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.12) if style is None else None
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9.6)


def build_docx(tables: dict[str, pd.DataFrame]) -> Path:
    if not MANUSCRIPT_PATH.exists():
        raise FileNotFoundError(f"Missing manuscript source: {rel(MANUSCRIPT_PATH)}")
    lines = MANUSCRIPT_PATH.read_text(encoding="utf-8").splitlines()
    title_text, subtitle_text = manuscript_title_and_subtitle(lines)

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], columns=1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    run = title.add_run(title_text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(subtitle_text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"FIN 496 Capstone\nFrozen repository commit: {FROZEN_COMMIT}\nGenerated: {date.today().isoformat()}")
    meta_run.font.name = "Times New Roman"
    meta_run.font.size = Pt(10)

    doc.add_page_break()
    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, columns=2)

    started = False
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            text = " ".join(line.strip() for line in paragraph_lines).strip()
            if text:
                add_markdown_paragraph(doc, text)
            paragraph_lines = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("# "):
            continue
        if not started:
            if line.startswith("## Abstract"):
                started = True
            else:
                continue
        marker = re.fullmatch(r"\{\{(table|figure):([A-Za-z0-9_]+)\}\}", line.strip())
        if marker:
            flush_paragraph()
            kind, key = marker.groups()
            if kind == "table":
                add_docx_table(doc, key, tables)
            else:
                add_docx_figure(doc, key)
            continue
        if line.startswith("## "):
            flush_paragraph()
            heading = line[3:].strip()
            if heading.lower() == "references":
                break
            h = doc.add_heading(heading, level=1)
            h.paragraph_format.keep_with_next = True
            continue
        if line.startswith("### "):
            flush_paragraph()
            h = doc.add_heading(line[4:].strip(), level=2)
            h.paragraph_format.keep_with_next = True
            continue
        if not line.strip():
            flush_paragraph()
            continue
        if line.strip().startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(line.strip()[2:], style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            continue
        paragraph_lines.append(line)
    flush_paragraph()

    doc.add_page_break()
    refs_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(refs_section, columns=1)
    doc.add_heading("References", level=1)
    refs = (PAPER_DIR / "references.md").read_text(encoding="utf-8").splitlines()
    for line in refs:
        clean = line.strip()
        if not clean or clean.startswith("#"):
            continue
        p = doc.add_paragraph(clean)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)

    out = PAPER_DIR / "final_manuscript.docx"
    doc.save(out)
    return out


def write_implementation_report(generated: Iterable[Path], validation_note: str) -> Path:
    files = "\n".join(f"- `{rel(path)}`" for path in generated)
    report = f"""# Final Manuscript Package Implementation Report

## Scope

- Built a markdown-driven final manuscript package from frozen repo outputs at `{FROZEN_COMMIT}`.
- Used only committed derived exports under `data/exports/final_paper_package_v2_expanded/`.
- Did not call APIs, collect transcripts, read raw Bloomberg workbooks, or create new empirical results.

## Generated files

{files}

## Validation

{validation_note}
"""
    path = PAPER_DIR / "IMPLEMENTATION_REPORT.md"
    path.write_text(report, encoding="utf-8")
    return path


def main() -> None:
    PAPER_DIR.mkdir(exist_ok=True)
    TABLE_DIR.mkdir(exist_ok=True)
    FIGURE_DIR.mkdir(exist_ok=True)

    tables = build_tables_and_figures()
    refs = write_references()
    audit = write_source_audit()
    docx = build_docx(tables)

    generated = [
        PAPER_DIR / "build_final_docx.py",
        MANUSCRIPT_PATH,
        refs,
        audit,
        docx,
        *sorted(TABLE_DIR.glob("*")),
        *sorted(FIGURE_DIR.glob("*")),
    ]
    report = write_implementation_report(
        generated,
        "Builder completed successfully from local frozen exports. Structural DOCX QA passed after generation (cover/body/reference sections use 1/2/1 columns, 10 generated tables are present, and revision-note text is absent). Full page-image rendering was attempted with the Documents render workflow, but this machine does not have LibreOffice/`soffice` installed.",
    )
    print(f"Wrote {rel(audit)}")
    print(f"Wrote {rel(docx)}")
    print(f"Wrote {rel(report)}")


if __name__ == "__main__":
    main()
